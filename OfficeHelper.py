# powered by: @御河DE天街
# 修改Excel文件的格式

class ExcelFormatHelper:

    # 加载Excel
    def __init__(self,file_path):
        import openpyxl
        self.file_path = file_path
        if self.file_path.split('.')[-1] == 'xlsx':
            self.wb = openpyxl.load_workbook(file_path)
        else:
            raise Exception('文件格式错误！')

    # 获取sheet页名称列表
    def get_sheet_names(self):
        sheet_names = [sheet_name for sheet_name in self.wb.sheetnames]
        return sheet_names


    # 更改某个单元格的字体格式
    def change_cell_format(self,sheet_name,row,column,font_format):
        self.wb[sheet_name].cell(row,column).font = font_format

    # 更改某个区域内单元格的字体格式
    def change_area_format(self,sheet_name,start_row,start_column,end_row,end_column,font_format):
        for i in range(start_row,end_row+1):
            for j in range(start_column,end_column+1):
                self.wb[sheet_name].cell(i,j).font = font_format

    # 更改某个sheet的所有字体格式 
    def change_font_format(self,sheet_name,font_format):
        for i in range(1,self.wb[sheet_name].max_row+1):
            for j in range(1,self.wb[sheet_name].max_column+1):
                self.wb[sheet_name].cell(i,j).font = font_format
    
    # 更改所有sheet的所有单元格字体格式
    def change_all_cell_format(self,cell_format):
        for sheet_name in self.wb.sheetnames:
            for i in range(1,self.wb[sheet_name].max_row+1):
                for j in range(1,self.wb[sheet_name].max_column+1):
                    self.wb[sheet_name].cell(i,j).font = cell_format
    

    # 更改某个单元格的数字格式
    def change_cell_number_format(self,sheet_name,row,column,number_format):
        if self.wb[sheet_name].cell(row,column) is int or self.wb[sheet_name].cell(row,column) is float:
            self.wb[sheet_name].cell(row,column).number_format = number_format
    
    # 更改某个区域内单元格的数字格式
    def change_area_number_format(self,sheet_name,start_row,start_column,end_row,end_column,number_format):
        for i in range(start_row,end_row+1):
            for j in range(start_column,end_column+1):
                if self.wb[sheet_name].cell(i,j) is int or self.wb[sheet_name].cell(i,j) is float:
                    self.wb[sheet_name].cell(i,j).number_format = number_format
    
    # 更改某个sheet的所有数字格式
    def change_number_format(self,sheet_name,number_format):
        for i in range(1,self.wb[sheet_name].max_row+1):
            for j in range(1,self.wb[sheet_name].max_column+1):
                if self.wb[sheet_name].cell(i,j) is int or self.wb[sheet_name].cell(i,j) is float:
                    self.wb[sheet_name].cell(i,j).number_format = number_format
    
    # 更改所有sheet的所有单元格数字格式
    def change_all_number_format(self,number_format):
        for sheet_name in self.wb.sheetnames:
            for i in range(1,self.wb[sheet_name].max_row+1):
                for j in range(1,self.wb[sheet_name].max_column+1):
                    if self.wb[sheet_name].cell(i,j) is int or self.wb[sheet_name].cell(i,j) is float:
                        self.wb[sheet_name].cell(i,j).number_format = number_format
    

    # 更改某个单元格的填充格式
    def change_cell_fill_format(self,sheet_name,row,column,fill_format):
        self.wb[sheet_name].cell(row,column).fill = fill_format
    
    # 更改某个区域内单元格的填充格式
    def change_area_fill_format(self,sheet_name,start_row,start_column,end_row,end_column,fill_format):
        for i in range(start_row,end_row+1):
            for j in range(start_column,end_column+1):
                self.wb[sheet_name].cell(i,j).fill = fill_format

    # 更改某个sheet的所有单元格填充格式
    def change_fill_format(self,sheet_name,fill_format):
        for i in range(1,self.wb[sheet_name].max_row+1):
            for j in range(1,self.wb[sheet_name].max_column+1):
                self.wb[sheet_name].cell(i,j).fill = fill_format
                
    # 更改所有sheet的所有单元格填充格式
    def change_all_fill_format(self,fill_format):
        for sheet_name in self.wb.sheetnames:
            for i in range(1,self.wb[sheet_name].max_row+1):
                for j in range(1,self.wb[sheet_name].max_column+1):
                    self.wb[sheet_name].cell(i,j).fill = fill_format


    # 更改某个单元格的对齐方式
    def change_cell_alignment(self,sheet_name,row,column,alignment):
        self.wb[sheet_name].cell(row,column).alignment = alignment
        
    # 更改某个区域内单元格的对齐方式
    def change_area_alignment(self,sheet_name,start_row,start_column,end_row,end_column,alignment):
        for i in range(start_row,end_row+1):
            for j in range(start_column,end_column+1):
                self.wb[sheet_name].cell(i,j).alignment = alignment
        
    # 调整某个sheet的所有单元格对齐方式
    def change_alignment(self,sheet_name,alignment):
        for i in range(1,self.wb[sheet_name].max_row+1):
            for j in range(1,self.wb[sheet_name].max_column+1):
                self.wb[sheet_name].cell(i,j).alignment = alignment
    
    # 调整所有sheet的所有单元格对齐方式
    def change_all_alignment(self,alignment):
        for sheet_name in self.wb.sheetnames:
            for i in range(1,self.wb[sheet_name].max_row+1):
                for j in range(1,self.wb[sheet_name].max_column+1):
                    self.wb[sheet_name].cell(i,j).alignment = alignment
    

    # 更改某个单元格的边框格式
    def change_cell_border(self,sheet_name,row,column,border):
        self.wb[sheet_name].cell(row,column).border = border
        
    # 更改某个区域内单元格的边框格式
    def change_area_border(self,sheet_name,start_row,start_column,end_row,end_column,border):
        for i in range(start_row,end_row+1):
            for j in range(start_column,end_column+1):
                self.wb[sheet_name].cell(i,j).border = border
                    
    # 调整某个sheet的所有单元格边框格式
    def change_border(self,sheet_name,border):
        for i in range(1,self.wb[sheet_name].max_row+1):
            for j in range(1,self.wb[sheet_name].max_column+1):
                self.wb[sheet_name].cell(i,j).border = border
                
    # 调整所有sheet的所有单元格边框格式
    def change_all_border(self,border):
        for sheet_name in self.wb.sheetnames:
            for i in range(1,self.wb[sheet_name].max_row+1):
                for j in range(1,self.wb[sheet_name].max_column+1):
                    self.wb[sheet_name].cell(i,j).border = border
    

    # 向某个单元格写入
    def write_to_cell(self,sheet_name,row,column,value):
        self.wb[sheet_name].cell(row,column).value = value
    
    # 向某个区域内写入
    def write_to_area(self,sheet_name,start_row,start_column,end_row,end_column,value):
        for i in range(start_row,end_row+1):
            for j in range(start_column,end_column+1):
                self.wb[sheet_name].cell(i,j).value = value
    

    # 删除某个单元格
    def delete_cell(self,sheet_name,row,column):
        self.wb[sheet_name].cell(row,column).value = None
        
    # 删除某个区域内的单元格
    def delete_area(self,sheet_name,start_row,start_column,end_row,end_column):
        for i in range(start_row,end_row+1):
            for j in range(start_column,end_column+1):
                self.wb[sheet_name].cell(i,j).value = None

    # 删除多列
    def delete_columns(self,sheet_name,start_column,end_column):
        for i in range(1,self.wb[sheet_name].max_row+1):
            for j in range(start_column,end_column+1):
                self.wb[sheet_name].cell(i,j).value = None
                
    # 删除多行
    def delete_rows(self,sheet_name,start_row,end_row):
        for i in range(start_row,end_row+1):
            for j in range(1,self.wb[sheet_name].max_column+1):
                self.wb[sheet_name].cell(i,j).value = None
    

    # 合并并居中多个单元格
    def merge_cells(self,sheet_name,start_row,start_column,end_row,end_column):
        from openpyxl.styles import Alignment
        self.wb[sheet_name].merge_cells(start_row=start_row,start_column=start_column,end_row=end_row,end_column=end_column)
        self.wb[sheet_name].cell(start_row,start_column).alignment = Alignment(horizontal='center',vertical='center')
    
    # 合并并居中多个区域内单元格
    def merge_area_cells(self,sheet_name,start_row,start_column,end_row,end_column):
        from openpyxl.styles import Alignment
        for i in range(start_row,end_row+1):
            for j in range(start_column,end_column+1):
                self.wb[sheet_name].cell(i,j).alignment = Alignment(horizontal='center',vertical='center')
        self.wb[sheet_name].merge_cells(start_row=start_row,start_column=start_column,end_row=end_row,end_column=end_column)
    
    # 合并多个sheet为一个sheet，并另存为另一个文件
    # 将sheet_name存入列表并排序即可实现按要求合并
    def merge_sheet(self,new_file_path):
        import pandas as pd
        data=[]
        for sheetname in self.wb.sheetnames:
            df1=pd.read_excel(self.file_path,sheet_name=sheetname,header=4)
            data.append(df1)
        df=pd.concat(data)
        df.to_excel(new_file_path,index=False)
    
    # 将指定的sheet合并为一个sheet并另存为另一个文件
    def merge_sheet_by_name(self,new_file_path,*sheet_name):
        import pandas as pd
        data=[]
        for sheetname in self.wb.sheetnames:
            if sheetname == sheet_name:
                df1=pd.read_excel(self.file_path,sheet_name=sheetname,header=4)
                data.append(df1)
        df=pd.concat(data)
        df.to_excel(new_file_path,index=False)


    # 获得某个值所在单元格的位置
    def get_cell_location(self,sheet_name,value):
        from openpyxl.utils import get_column_letter
        for i in range(1,self.wb[sheet_name].max_row+1):
            for j in range(1,self.wb[sheet_name].max_column+1):
                if self.wb[sheet_name].cell(i,j).value == value:
                    return [get_column_letter(j),i],[i,j]
    
    # 查找某个值并进行替换
    def find_and_replace(self,sheet_name,find_value,replace_value):
        for i in range(1,self.wb[sheet_name].max_row+1):
            for j in range(1,self.wb[sheet_name].max_column+1):
                if self.wb[sheet_name].cell(i,j).value == find_value:
                    self.wb[sheet_name].cell(i,j).value = replace_value
    
    # 查找某个值所在的sheet
    def find_sheet_by_value(self,value):
        from openpyxl.utils import get_column_letter
        for sheet_name in self.wb.sheetnames:
            for i in range(1,self.wb[sheet_name].max_row+1):
                for j in range(1,self.wb[sheet_name].max_column+1):
                    if self.wb[sheet_name].cell(i,j).value == value:
                        return sheet_name,[get_column_letter(j),i],[i,j]
    
    

    # 设置列宽
    def set_column_width(self,sheet_name,column,width):
        self.wb[sheet_name].column_dimensions[column].width = width
        
    # 设置行高
    def set_row_height(self,sheet_name,row,height):
        self.wb[sheet_name].row_dimensions[row].height = height

    # 自适应列宽
    def auto_fit_column(self,sheet_name):
        from openpyxl.utils import get_column_letter
        lks=[]
        for i in range(1,self.wb[sheet_name].max_column+1):
            lk=1
            for j in range(1,self.wb[sheet_name].max_row+1):
                sz=self.wb[sheet_name].cell(row=j,column=i).value
                if isinstance(sz,str):
                    lk1=len(sz.encode('gbk'))
                else:
                    lk1=len(str(sz))
                if lk<lk1:
                    lk=lk1
            lks.append(lk)

        for i in range(1,self.wb[sheet_name].max_column+1):
            k=get_column_letter(i)
            self.wb[sheet_name].column_dimensions[k].width=lks[i-1]+2
    
    # 所有sheet自适应列宽
    def auto_fit_all_columns(self):
        for sheet_name in self.wb.sheetnames:
            self.auto_fit_column(sheet_name)


    # 关闭Excel文件
    def close_excel(self):
        self.wb.close()
        print('Excel文件关闭成功!')
    
    # 修改sheet名称
    def change_sheet_name(self,sheet_name,new_sheet_name):
        self.wb[sheet_name].title = new_sheet_name
    
    # 删除某个sheet
    def delete_sheet(self,sheet_name):
        self.wb.remove(self.wb[sheet_name])
    
    # 删除所有sheet
    def delete_all_sheet(self):
        for sheet_name in self.wb.sheetnames:
            self.wb.remove(self.wb[sheet_name])
    
    # 新建sheet
    def new_sheet(self,sheet_name):
        self.wb.create_sheet(sheet_name)
    
    # 修改文件名称
    def change_file_name(self,new_file_name):
        self.wb.save(new_file_name)
    
    # 修改文件类型
    def change_file_type(self,new_file_type):
        self.wb.save(self.file_path.split('.')[0]+'.'+new_file_type)

    # 删除Excel文件
    def delete_excel(self):
        import os
        if self.file_path.split('.')[-1] == 'xlsx':
            os.remove(self.file_path)

    # 保存Excel文件
    def save_excel(self):
        self.wb.save(self.file_path)
        print('Excel文件保存成功!')



# 修改word文件的格式
class WordFormatHelper:
    
    # 加载Word
    def __init__(self, word_path):
        from docx import Document
        self.word_path = word_path
        if self.word_path.split('.')[-1] == 'docx':
            self.doc = Document(word_path)
        else:
            raise Exception('文件格式错误!')
    
    # 读取Word文件内的全部文字内容
    def read_word(self):
        Text = []
        for para in self.doc.paragraphs:
            Text.append(para.text)
        return '\n'.join(Text)
    
    # 修改字体格式
    def change_format(self, font_name,font_size, font_color):
        from docx.shared import RGBColor
        from docx.shared import Pt
        from docx.oxml.ns import qn
        self.doc.styles['Normal'].font.size = Pt(font_size)
        self.doc.styles['Normal'].font.color.rgb = RGBColor(*font_color)
        for i in range(len(self.doc.paragraphs)): 
            for run in self.doc.paragraphs[i].runs:
                run.font.name = font_name
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                # run.font.bold = False
                # run.font.italic = False
                # run.font.underline = False
                # run.font.strike = False
                # run.font.shadow = False
                # run.font.highlight_color = None
    
    # 修改某个段落的字体格式
    def change_format_of_paragraph(self, paragraph_index, font_name,font_size, font_color):
        from docx.shared import RGBColor
        from docx.shared import Pt
        from docx.oxml.ns import qn
        self.doc.paragraphs[paragraph_index].style.font.size = Pt(font_size)
        self.doc.paragraphs[paragraph_index].style.font.color.rgb = RGBColor(*font_color)
        for run in self.doc.paragraphs[paragraph_index].runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            # run.font.bold = False
            # run.font.italic = False
            # run.font.underline = False
            # run.font.strike = False
            # run.font.shadow = False
            # run.font.highlight_color = None

    # 修改对齐方式
    def change_align(self, align):
        for i in range(len(self.doc.paragraphs)):
            self.doc.paragraphs[i].alignment = align
    
    # 修改某个段落的对齐方式
    def change_align_of_paragraph(self, paragraph_index, align):
        self.doc.paragraphs[paragraph_index].alignment = align

    # 设置行间距
    def change_line_spacing(self, line_spacing):
        for i in range(len(self.doc.paragraphs)):
            self.doc.paragraphs[i].paragraph_format.line_spacing_rule = line_spacing
    
    # 设置某个段落的行间距
    def change_line_spacing_of_paragraph(self, paragraph_index, line_spacing):
        from docx.enum.text import WD_LINE_SPACING
        self.doc.paragraphs[paragraph_index].paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        self.doc.paragraphs[paragraph_index].paragraph_format.line_spacing = line_spacing
    
    # 设置段落间距
    def change_paragraph_spacing(self, paragraph_spacing1,paragraph_spacing2):
        for i in range(len(self.doc.paragraphs)):
            self.doc.paragraphs[i].paragraph_format.space_after = paragraph_spacing1
            self.doc.paragraphs[i].paragraph_format.space_before = paragraph_spacing2
    
    # 设置某个段落的段落间距
    def change_paragraph_spacing_of_paragraph(self, paragraph_index, paragraph_spacing1,paragraph_spacing2):
        self.doc.paragraphs[paragraph_index].paragraph_format.space_after = paragraph_spacing1
        self.doc.paragraphs[paragraph_index].paragraph_format.space_before = paragraph_spacing2
    
    # 首行缩进2个字符
    def indent_first_line(self):
        for i in range(len(self.doc.paragraphs)):
            self.doc.paragraphs[i].paragraph_format.first_line_indent = self.doc.paragraphs[i].style.font.size*2
    
    # 某个段落首行缩进2个字符
    def indent_first_line_of_paragraph(self, paragraph_index):
        self.doc.paragraphs[paragraph_index].paragraph_format.first_line_indent = self.doc.paragraphs[paragraph_index].style.font.size*2

    # 取消首行缩进
    def unindent_first_line(self):
        for i in range(len(self.doc.paragraphs)):
            self.doc.paragraphs[i].paragraph_format.first_line_indent = 0
    
    # 取消某个段落首行缩进
    def unindent_first_line_of_paragraph(self, paragraph_index):
        self.doc.paragraphs[paragraph_index].paragraph_format.first_line_indent = 0
    
    # 批量替换
    def replace_all(self, old_str, new_str):
        for paragraghs in self.doc.paragraphs:
            for run in paragraghs.runs:
                run.text = run.text.replace(old_str, new_str)
    
    # 某个段落批量替换
    def replace_all_of_paragraph(self, paragraph_index, old_str, new_str):
        for run in self.doc.paragraphs[paragraph_index].runs:
            run.text = run.text.replace(old_str, new_str)

    # 删除Word文件
    def delete_word(self):
        import os
        os.remove(self.word_path)
    
    # 修改Word文件名
    def rename_word(self, new_name):
        import os
        os.rename(self.word_path, new_name)
    
    # 保存Word文件
    def save_word(self, save_path):
        self.doc.save(save_path)
        print('Word文件保存成功!')
